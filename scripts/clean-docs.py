#!/usr/bin/env python3
"""
clean-docs.py

Scrubs "Cenomi/سينومي/Seinomy" and "Bot/بوت" from marketing .docx files.
Produces *_clean.docx next to the originals.

Requires: python-docx
    pip install python-docx
"""

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    sys.stderr.write("Please install python-docx:  pip install python-docx\n")
    sys.exit(1)

# Replacement map. Order matters: longer phrases first to avoid partial matches.
REPLACEMENTS = [
    ("WhatsApp Bot",      "المساعد الذكي"),
    ("Telegram Bot",      "المساعد الذكي"),
    ("بوت تيليجرام",     "المساعد الذكي"),
    ("بوت واتساب",       "المساعد الذكي"),
    ("البوت",            "المساعد الذكي"),
    ("Cenomi API",       "API التكامل"),
    ("Cenomi",           "منصة التكامل"),
    ("CENOMI",           "منصة التكامل"),
    ("Seinomy",          "منصة التكامل"),
    ("seinomy",          "منصة التكامل"),
    ("سينومي",           "منصة التكامل"),
    ("بوت",              "المساعد الذكي"),
    ("Bot",              "Assistant"),
]


def replace_in_runs(runs):
    """Replace text in a list of runs, preserving formatting as much as possible."""
    for run in runs:
        text = run.text
        if not text:
            continue
        new_text = text
        for old, new in REPLACEMENTS:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != text:
            run.text = new_text


def process_paragraph(p):
    """
    Replacements across a paragraph's runs can lose formatting when a match
    spans multiple runs. We apply both per-run replacement and a full-text
    fallback for split matches.
    """
    replace_in_runs(p.runs)

    # Fallback: if the whole paragraph text still contains any target phrase
    # (because it was split across runs), rewrite all text into the first run.
    full = "".join(r.text for r in p.runs)
    replaced = full
    for old, new in REPLACEMENTS:
        if old in replaced:
            replaced = replaced.replace(old, new)
    if replaced != full and p.runs:
        # Clear every run, then stuff everything into the first.
        for r in p.runs[1:]:
            r.text = ""
        p.runs[0].text = replaced


def process_document(doc):
    # Body paragraphs
    for p in doc.paragraphs:
        process_paragraph(p)

    # Tables (may contain nested paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

    # Headers and footers
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                process_paragraph(p)


def scrub(infile: Path) -> Path:
    outfile = infile.with_name(infile.stem + "_clean" + infile.suffix)
    doc = Document(str(infile))
    process_document(doc)
    doc.save(str(outfile))
    return outfile


def main():
    files = [
        Path(__file__).parent.parent / "Musharaka Sales Management System .docx",
        Path(__file__).parent.parent / "Musharaka_Development_Plan.docx",
    ]

    for f in files:
        if not f.exists():
            print(f"[SKIP] Not found: {f.name}")
            continue
        try:
            out = scrub(f)
            print(f"[OK]   Wrote: {out.name}")
        except Exception as e:
            print(f"[FAIL] {f.name}: {e}")


if __name__ == "__main__":
    main()
